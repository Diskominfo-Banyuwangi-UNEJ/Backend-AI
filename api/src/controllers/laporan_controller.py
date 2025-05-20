from flask import Blueprint, jsonify, request, send_file, make_response, render_template_string
from src.models.laporan_model import Laporan, StatusPengerjaan, Kategori
from src.models.user_model import User
from src import db
from sqlalchemy.exc import IntegrityError
from datetime import datetime
from functools import wraps
from io import BytesIO
from reportlab.pdfgen import canvas
from docx import Document
from xhtml2pdf import pisa
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

laporan = Blueprint('laporan', __name__)

def handle_errors(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        try:
            return f(*args, **kwargs)
        except IntegrityError:
            db.session.rollback()
            return jsonify({
                'status': "error",
                'message': "Database integrity error occurred"
            }), 400
        except Exception as e:
            db.session.rollback()
            return jsonify({
                'status': "error",
                'message': str(e)
            }), 500
    return decorated_function

@laporan.route('/getAllLaporan', methods=["GET"])
@handle_errors
def get_all_laporan():
    try:
        status_filter = request.args.get('status')
        kategori_filter = request.args.get('kategori')
        
        query = Laporan.query
        
        # Apply filters if they exist
        if status_filter:
            query = query.filter_by(status_pengerjaan=StatusPengerjaan(status_filter))
        if kategori_filter:
            query = query.filter_by(kategori=Kategori(kategori_filter))
        
        laporan_list = query.order_by(Laporan.created_at.desc()).all()
        
        if not laporan_list:
            return jsonify({
                'status': "success",
                'message': "Tidak ada laporan ditemukan.",
                'data': []
            }), 200
            
        return jsonify({
            'status': "success",
            'data': [laporan.to_dict() for laporan in laporan_list]
        }), 200
    except ValueError as e:
        return jsonify({
            'status': "error",
            'message': "Status atau kategori tidak valid"
        }), 400
    except Exception as e:
        return jsonify({
            'status': "error",
            'message': str(e)
        }), 500

@laporan.route('/getLaporanDetail/<int:id>', methods=["GET"])
@handle_errors
def get_laporan_detail(id):
    try:
        laporan = Laporan.query.get(id)
        if not laporan:
            return jsonify({
                'status': "error",
                'message': "Laporan tidak ditemukan."
            }), 404
            
        return jsonify({
            'status': "success",
            'data': laporan.to_dict()
        }), 200
    except Exception as e:
        return jsonify({
            'status': "error",
            'message': str(e)
        }), 500

@laporan.route('/createLaporan', methods=["POST"])
@handle_errors
def create_laporan():
    try:
        data = request.get_json()
        
        # Validate required fields
        required_fields = ['judul_laporan', 'deskripsi', 'estimasi', 'kategori', 'id_user']
        for field in required_fields:
            if field not in data or not data[field]:
                return jsonify({
                    'status': "error",
                    'message': f"Field {field} harus diisi!"
                }), 400
        
        try:
            kategori = Kategori(data['kategori'])
        except ValueError:
            return jsonify({
                'status': "error",
                'message': "Kategori tidak valid. Pilih antara KERAMAIAN atau TUMPUKAN_SAMPAH"
            }), 400
                
        new_laporan = Laporan(
            judul_laporan=data['judul_laporan'],
            deskripsi=data['deskripsi'],
            estimasi=data['estimasi'],
            kategori=kategori,
            status_pengerjaan=StatusPengerjaan.BARU,
            id_user=data['id_user']
        )
        
        db.session.add(new_laporan)
        db.session.commit()
        
        return jsonify({
            'status': "success",
            'message': "Laporan berhasil dibuat.",
            'data': new_laporan.to_dict()
        }), 201
    except Exception as e:
        return jsonify({
            'status': "error",
            'message': str(e)
        }), 500

@laporan.route('/updateLaporan/<int:id>', methods=["PUT"])
@handle_errors
def update_laporan(id):
    try:
        laporan = Laporan.query.get(id)
        if not laporan:
            return jsonify({
                'status': "error",
                'message': "Laporan tidak ditemukan."
            }), 404
            
        data = request.get_json()
        
        # Update fields if they exist in request
        if 'judul_laporan' in data:
            laporan.judul_laporan = data['judul_laporan']
        if 'deskripsi' in data:
            laporan.deskripsi = data['deskripsi']
        if 'estimasi' in data:
            laporan.estimasi = data['estimasi']
        if 'kategori' in data:
            try:
                laporan.kategori = Kategori(data['kategori'])
            except ValueError:
                return jsonify({
                    'status': "error",
                    'message': "Kategori tidak valid. Pilih antara KERAMAIAN atau TUMPUKAN_SAMPAH"
                }), 400
        if 'status_pengerjaan' in data:
            try:
                laporan.status_pengerjaan = StatusPengerjaan(data['status_pengerjaan'])
            except ValueError:
                return jsonify({
                    'status': "error",
                    'message': "Status tidak valid. Pilih antara BARU, DIBACA, atau SELESAI"
                }), 400
            
        db.session.commit()
        
        return jsonify({
            'status': "success",
            'message': "Laporan berhasil diperbarui.",
            'data': laporan.to_dict()
        }), 200
    except Exception as e:
        return jsonify({
            'status': "error",
            'message': str(e)
        }), 500

@laporan.route('/deleteLaporan/<int:id>', methods=["DELETE"])
@handle_errors
def delete_laporan(id):
    try:
        laporan = Laporan.query.get(id)
        if not laporan:
            return jsonify({
                'status': "error",
                'message': "Laporan tidak ditemukan."
            }), 404
            
        db.session.delete(laporan)
        db.session.commit()
        
        return jsonify({
            'status': "success",
            'message': "Laporan berhasil dihapus."
        }), 200
    except Exception as e:
        return jsonify({
            'status': "error",
            'message': str(e)
        }), 500


@laporan.route('/updateStatusLaporan/<int:id>', methods=["PATCH"])
@handle_errors
def update_status_laporan(id):
    """Endpoint khusus untuk mengubah status laporan"""
    try:
        laporan = Laporan.query.get(id)
        if not laporan:
            return jsonify({
                'status': "error",
                'message': "Laporan tidak ditemukan."
            }), 404
            
        data = request.get_json()
        
        if 'status_pengerjaan' not in data:
            return jsonify({
                'status': "error",
                'message': "Field status_pengerjaan harus diisi!"
            }), 400
            
        try:
            new_status = StatusPengerjaan(data['status_pengerjaan'])
            laporan.status_pengerjaan = new_status
            db.session.commit()
            
            return jsonify({
                'status': "success",
                'message': "Status berhasil diperbarui.",
                'data': {
                    'id': laporan.id,
                    'status_pengerjaan': laporan.status_pengerjaan.value
                }
            }), 200
        except ValueError:
            return jsonify({
                'status': "error",
                'message': "Status tidak valid. Pilih antara BARU, DIBACA, atau SELESAI"
            }), 400
    except Exception as e:
        return jsonify({
            'status': "error",
            'message': str(e)
        }), 500


def get_status_class(status):
    status_map = {
        'BARU': 'new',
        'DIBACA': 'progress',
        'SELESAI': 'completed'
    }
    return status_map.get(status, 'new')

@laporan.route('/downloadLaporan/<int:id>/<format>', methods=["GET"])
@handle_errors
def download_laporan(id, format):
    try:
        laporan = Laporan.query.get(id)
        if not laporan:
            return jsonify({
                'status': "error",
                'message': "Laporan tidak ditemukan."
            }), 404
            
        if format not in ['pdf', 'docx']:
            return jsonify({
                'status': "error",
                'message': "Format tidak valid. Pilih antara pdf atau docx"
            }), 400

        # Prepare data for template
        report_data = {
            'report_id': laporan.id,
            'judul_laporan': laporan.judul_laporan,
            'deskripsi': laporan.deskripsi,
            'kategori': laporan.kategori.value,
            'status': laporan.status_pengerjaan.value,
            'status_class': get_status_class(laporan.status_pengerjaan.value),
            'created_at': laporan.created_at.strftime('%d %B %Y %H:%M'),
            'print_date': datetime.now().strftime('%d %B %Y %H:%M'),
            'printed_by': 'Sistem Pelaporan'
        }

        if format == 'pdf':
            # Render HTML template
            with open('src/templates/report_template.html', 'r') as f:
                template = f.read()
            html = render_template_string(template, **report_data)
            
            # Create PDF
            buffer = BytesIO()
            pisa.CreatePDF(html, dest=buffer, encoding='UTF-8')
            
            buffer.seek(0)
            response = make_response(buffer.getvalue())
            response.headers['Content-Type'] = 'application/pdf'
            response.headers['Content-Disposition'] = (
                f'attachment; filename=Laporan_{laporan.id}_{datetime.now().strftime("%Y%m%d")}.pdf'
            )
            return response
            
        elif format == 'docx':
            # Create DOCX with styling
            document = Document()
            
            # Add header
            header = document.sections[0].header
            header_para = header.paragraphs[0]
            header_run = header_para.add_run()
            # header_run.add_picture('static/images/logo.png', width=Inches(1))
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            title = document.add_heading('LAPORAN RESMI', level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            document.add_paragraph('Sistem Pelaporan Masyarakat', style='Subtitle')
            document.add_paragraph()
            
            # Add report info
            document.add_heading('INFORMASI UMUM', level=2)
            info_table = document.add_table(rows=4, cols=2)
            info_table.style = 'Light Shading'
            
            cells = info_table.rows[0].cells
            cells[0].text = 'Nomor Laporan'
            cells[1].text = f': {laporan.id}'
            
            cells = info_table.rows[1].cells
            cells[0].text = 'Tanggal'
            cells[1].text = f': {laporan.created_at.strftime("%d %B %Y %H:%M")}'
            
            cells = info_table.rows[2].cells
            cells[0].text = 'Status'
            status_run = cells[1].paragraphs[0].add_run(f': {laporan.status_pengerjaan.value}')
            status_run.bold = True
            if laporan.status_pengerjaan.value == 'SELESAI':
                status_run.font.color.rgb = RGBColor(0x4C, 0xAF, 0x50)
            elif laporan.status_pengerjaan.value == 'DIBACA':
                status_run.font.color.rgb = RGBColor(0x21, 0x96, 0xF3)
            else:
                status_run.font.color.rgb = RGBColor(0xFF, 0xC1, 0x07)
            
            cells = info_table.rows[3].cells
            cells[0].text = 'Jenis Laporan'
            cells[1].text = f': {laporan.kategori.value}'
            
            # Add report content
            document.add_heading('DETAIL LAPORAN', level=2)
            document.add_heading(laporan.judul_laporan, level=3)
            document.add_paragraph(laporan.deskripsi)
            
            # Add footer
            footer = document.sections[0].footer
            footer_para = footer.paragraphs[0]
            footer_para.text = (
                f"Dokumen ini dicetak pada {datetime.now().strftime('%d %B %Y %H:%M')} oleh Sistem Pelaporan\n"
                "© 2023 Sistem Pelaporan Masyarakat"
            )
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_para.style.font.size = Pt(10)
            
            buffer = BytesIO()
            document.save(buffer)
            buffer.seek(0)
            
            response = make_response(buffer.getvalue())
            response.headers['Content-Type'] = (
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response.headers['Content-Disposition'] = (
                f'attachment; filename=Laporan_{laporan.id}_{datetime.now().strftime("%Y%m%d")}.docx'
            )
            return response
            
    except Exception as e:
        return jsonify({
            'status': "error",
            'message': f"Gagal mengunduh laporan. Silakan coba lagi. Error: {str(e)}"
        }), 500
    """Endpoint untuk mengunduh laporan dalam format PDF atau DOCX"""
    try:
        laporan = Laporan.query.get(id)
        if not laporan:
            return jsonify({
                'status': "error",
                'message': "Laporan tidak ditemukan."
            }), 404
            
        if format not in ['pdf', 'docx']:
            return jsonify({
                'status': "error",
                'message': "Format tidak valid. Pilih antara pdf atau docx"
            }), 400
            
        if format == 'pdf':
            # Create PDF
            buffer = BytesIO()
            p = canvas.Canvas(buffer)
            
            # Add document content
            p.drawString(100, 800, f"Judul Laporan: {laporan.judul_laporan}")
            p.drawString(100, 780, f"Jenis Laporan: {laporan.kategori.value}")
            p.drawString(100, 760, f"Status: {laporan.status_pengerjaan.value}")
            p.drawString(100, 740, f"Tanggal: {laporan.created_at.strftime('%Y-%m-%d %H:%M')}")
            p.drawString(100, 700, "Deskripsi:")
            text = p.beginText(100, 680)
            text.textLines(laporan.deskripsi)
            p.drawText(text)
            
            p.showPage()
            p.save()
            
            buffer.seek(0)
            response = make_response(buffer.getvalue())
            response.headers['Content-Type'] = 'application/pdf'
            response.headers['Content-Disposition'] = f'attachment; filename=laporan_{id}.pdf'
            return response
            
        elif format == 'docx':
            # Create DOCX
            document = Document()
            document.add_heading(f'Laporan #{laporan.id}', 0)
            
            document.add_paragraph(f"Judul Laporan: {laporan.judul_laporan}")
            document.add_paragraph(f"Jenis Laporan: {laporan.kategori.value}")
            document.add_paragraph(f"Status: {laporan.status_pengerjaan.value}")
            document.add_paragraph(f"Tanggal: {laporan.created_at.strftime('%Y-%m-%d %H:%M')}")
            document.add_paragraph("Deskripsi:")
            document.add_paragraph(laporan.deskripsi)
            
            buffer = BytesIO()
            document.save(buffer)
            buffer.seek(0)
            
            response = make_response(buffer.getvalue())
            response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            response.headers['Content-Disposition'] = f'attachment; filename=laporan_{id}.docx'
            return response
            
    except Exception as e:
        return jsonify({
            'status': "error",
            'message': f"Gagal mengunduh laporan. Silakan coba lagi. Error: {str(e)}"
        }), 500